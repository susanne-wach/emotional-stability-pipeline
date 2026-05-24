"""
Premium-PDF-Builder für die Emotional Stability Analyse.

Liest ein Auswertungs-JSON (aus output/) und rendert ein .docx im
Your-Balance-Dokumenten-Brand. Konvertierung nach PDF passiert separat
(LibreOffice / docx2pdf — Phase 2).

Output: ~12-15 Seiten branded Word-Dokument.
"""

from __future__ import annotations

import json
import sys
from datetime import datetime
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT / "src"))

from your_balance_docx import (
    create_branded_doc, BRAND,
    add_brand_cover, add_brand_h1, add_brand_h2,
    add_brand_paragraph, add_brand_italic_note,
    add_brand_bullet, add_brand_number,
    add_brand_pagebreak, add_brand_ornament,
    add_brand_inline_caps_label, add_brand_callout,
    add_brand_quote,
    save_branded_doc,
)

# Logo: liegt im Public Repo unter assets/logo.png
LOGO_PATH = PROJECT_ROOT / "assets" / "logo.png"


def _add_hyperlink(paragraph, url, text, font_name="Georgia", font_size=9, color="888888"):
    """Fügt einen anklickbaren Hyperlink in einen Paragraph ein."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size * 2))  # halbe Punkte
    rPr.append(sz)

    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)

    # Underline für Hyperlink-Look
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'none')  # KEIN Underline, dezent
    rPr.append(u)

    new_run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_header_footer(doc):
    """
    Fügt auf jeder Seite (außer Cover) ein:
      - HEADER: kleines YB-Logo zentriert + Abstand darunter
      - FOOTER: your-balance.at · susanne@your-balance.at (anklickbar) | Seite X
    """
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    section = doc.sections[0]

    # ---- HEADER ----
    header = section.header
    header.is_linked_to_previous = False
    for p in list(header.paragraphs):
        p.clear()

    # Logo zentriert
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run()
    if LOGO_PATH.exists():
        run.add_picture(str(LOGO_PATH), width=Cm(2.6))  # klein, dezent

    # Abstand unter dem Logo (zwei leere Paragraphen)
    hp.paragraph_format.space_after = Pt(12)
    spacer = header.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer2 = header.add_paragraph()
    spacer2.paragraph_format.space_after = Pt(0)

    # ---- FOOTER ----
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        p.clear()
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Anklickbarer Website-Link
    _add_hyperlink(fp, "https://your-balance.at", "your-balance.at")

    # Trenner
    sep1 = fp.add_run("  ·  ")
    sep1.font.name = "Georgia"
    sep1.font.size = Pt(9)
    sep1.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Anklickbarer Mail-Link
    _add_hyperlink(fp, "mailto:susanne@your-balance.at", "susanne@your-balance.at")

    # Trenner + Seite
    sep2 = fp.add_run("  ·  Seite ")
    sep2.font.name = "Georgia"
    sep2.font.size = Pt(9)
    sep2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # PAGE-Feld einfügen
    page_run = fp.add_run()
    page_run.font.name = "Georgia"
    page_run.font.size = Pt(9)
    page_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    page_run._r.append(fldChar1)
    page_run._r.append(instrText)
    page_run._r.append(fldChar2)

    # Auf Cover-Seite Header/Footer NICHT zeigen
    sectPr = section._sectPr
    titlePg = sectPr.find(qn('w:titlePg'))
    if titlePg is None:
        titlePg = OxmlElement('w:titlePg')
        sectPr.append(titlePg)


def _ersten_namen(name: str) -> str:
    """„Antonia (fiktiv)" → „Antonia"."""
    return name.split("(")[0].split()[0].strip()


def _absaetze(text: str):
    """Yields jeden Absatz separat (split bei \\n\\n)."""
    for absatz in text.split("\n\n"):
        absatz = absatz.strip()
        if absatz:
            yield absatz


def _balken(score: float, breite: int = 30) -> str:
    """Textuelle Score-Visualisierung — z.B. ████████░░░░░░░░ 64%."""
    gefuellt = int(round(score / 100 * breite))
    leer = breite - gefuellt
    return "█" * gefuellt + "░" * leer


def render_auswertung_pdf(auswertung_json_path: Path, output_path: Path):
    data = json.loads(Path(auswertung_json_path).read_text(encoding="utf-8"))
    name_voll = data["meta"]["name"]
    name_kurz = _ersten_namen(name_voll)
    scores = data["scores"]
    typen = data["typen"]
    a = data["auswertung"]  # die Premium-Texte

    doc = create_branded_doc()

    # ---------- COVER ----------
    add_brand_cover(
        doc,
        caps_label="PREMIUM-AUSWERTUNG",
        title="Emotional Stability\nAnalyse™",
        subtitle=f"Persönliche Tiefenanalyse für {name_kurz}",
        attribution="Susanne Wachter — Your Balance",
        url="your-balance.at",
        date_text=datetime.now().strftime("Stand: %d. %B %Y").replace("January","Januar").replace("February","Februar").replace("March","März").replace("April","April").replace("May","Mai").replace("June","Juni").replace("July","Juli").replace("August","August").replace("September","September").replace("October","Oktober").replace("November","November").replace("December","Dezember"),
    )
    add_brand_pagebreak(doc)

    # ---------- TEIL 1 — SPIEGELUNG ----------
    add_brand_h1(doc, "Bevor wir starten")
    add_brand_paragraph(doc, a["anrede"])
    for absatz in _absaetze(a["teil_1_spiegelung"]):
        add_brand_paragraph(doc, absatz)
    add_brand_ornament(doc)
    add_brand_pagebreak(doc)

    # ---------- TEIL 2 — SCORES ----------
    add_brand_h1(doc, "Deine Werte")
    add_brand_paragraph(doc, a["teil_2_scores"]["kurz_einleitung"])
    add_brand_paragraph(doc, "")

    # ESI als Hervorhebung
    add_brand_inline_caps_label(doc, "EMOTIONAL STABILITY INDEX™")
    esi = scores["emotional_stability_index"]
    add_brand_callout(doc, f"{esi:.1f} / 100")
    add_brand_paragraph(doc, "")
    # NEU: emotionale Score-Interpretation (dichte, poetische Sprache)
    add_brand_inline_caps_label(doc, "WAS DAS BEDEUTET")
    esi_text = a["teil_2_scores"].get("esi_was_das_bedeutet") or a["teil_2_scores"].get("esi_interpretation", "")
    for absatz in _absaetze(esi_text):
        add_brand_paragraph(doc, absatz)
    add_brand_paragraph(doc, "")

    # Vier Tiefenscores als Tabellen-artige Liste
    add_brand_h2(doc, "Deine Tiefenscores")
    interp = a["teil_2_scores"]["tiefenscores_interpretationen"]
    for key, label in [
        ("emotional_activation", "Emotionaler Aktivierungsgrad"),
        ("nervous_system_load", "Nervensystem-Belastung"),
        ("selbstverlust", "Selbstverlust-Level"),
        ("hoffnungsschleifen", "Hoffnungsschleifen-Level"),
    ]:
        score = scores[key]
        add_brand_inline_caps_label(doc, f"{label.upper()} — {score:.0f}%")
        add_brand_italic_note(doc, _balken(score))
        for absatz in _absaetze(interp.get(key, "")):
            add_brand_paragraph(doc, absatz)
        add_brand_paragraph(doc, "")
    add_brand_pagebreak(doc)

    # ---------- TEIL 3 — HAUPTMUSTER ----------
    hm = a["teil_3_hauptmuster"]
    add_brand_h1(doc, "Dein Hauptmuster")
    add_brand_h2(doc, hm["name"])
    add_brand_inline_caps_label(doc, f"DOMINANZ: {typen['haupttyp_score']:.0f}%")
    add_brand_paragraph(doc, "")
    add_brand_inline_caps_label(doc, "WAS WIR SEHEN")
    for absatz in _absaetze(hm["was_wir_sehen"]):
        add_brand_paragraph(doc, absatz)
    add_brand_paragraph(doc, "")

    # NEU: WOW-Moment (Susannes Premium-Hebel) — falls vorhanden
    heimlich = hm.get("was_du_vermutlich_heimlich_kennst")
    if heimlich:
        add_brand_inline_caps_label(doc, "WAS DU VERMUTLICH HEIMLICH KENNST")
        for absatz in _absaetze(heimlich):
            add_brand_paragraph(doc, absatz)
        add_brand_paragraph(doc, "")

    add_brand_inline_caps_label(doc, "WAS DICH WIRKLICH FESTHÄLT")
    for absatz in _absaetze(hm["was_dich_festhaelt"]):
        add_brand_paragraph(doc, absatz)
    add_brand_paragraph(doc, "")

    # NEU: "Was wir besonders deutlich sehen" — 3 Premium-Bulletpoints
    besonders = a.get("teil_3_5_was_wir_besonders_sehen")
    if besonders:
        add_brand_ornament(doc)
        add_brand_paragraph(doc, "")
        add_brand_h2(doc, "Was wir besonders deutlich erkennen")
        for punkt in besonders:
            add_brand_bullet(doc, punkt)
        add_brand_paragraph(doc, "")

    add_brand_ornament(doc)
    add_brand_pagebreak(doc)

    # ---------- TEIL 4 — ZWEITMUSTER (optional) ----------
    zm = a["teil_4_zweitmuster"]
    if zm.get("vorhanden"):
        add_brand_h1(doc, "Dein zweites Muster")
        add_brand_h2(doc, zm["name"])
        if typen.get("nebentyp_score"):
            add_brand_inline_caps_label(doc, f"PRÄSENZ: {typen['nebentyp_score']:.0f}%")
            add_brand_paragraph(doc, "")
        for absatz in _absaetze(zm["was_zusaetzlich_sichtbar_wird"]):
            add_brand_paragraph(doc, absatz)
        add_brand_ornament(doc)
        add_brand_pagebreak(doc)

    # ---------- TEIL 5 — BLINDER FLECK ----------
    bf = a["teil_5_blinder_fleck"]
    add_brand_h1(doc, "Dein blinder Fleck")
    add_brand_inline_caps_label(doc, "WAS DU VERSUCHST")
    add_brand_paragraph(doc, bf["was_du_versuchst"])
    add_brand_paragraph(doc, "")
    add_brand_inline_caps_label(doc, "WAS DEIN SYSTEM EIGENTLICH BRAUCHT")
    add_brand_paragraph(doc, bf["was_dein_system_eigentlich_braucht"])
    add_brand_paragraph(doc, "")
    add_brand_inline_caps_label(doc, "WARUM DAS SO IST")
    for absatz in _absaetze(bf["ausformulierung"]):
        add_brand_paragraph(doc, absatz)
    add_brand_ornament(doc)
    add_brand_pagebreak(doc)

    # ---------- TEIL 5.5 — WAS SICH ZUERST VERÄNDERN DARF (optional) ----------
    zuerst = a.get("teil_5_5_was_sich_zuerst_veraendern_darf")
    if zuerst:
        add_brand_h1(doc, "Was sich zuerst verändern darf")
        for absatz in _absaetze(zuerst):
            add_brand_paragraph(doc, absatz)
        add_brand_ornament(doc)
        add_brand_pagebreak(doc)

    # ---------- TEIL 6 — MINI-ÜBUNG ----------
    mu = a["teil_6_mini_uebung"]
    add_brand_h1(doc, "Deine erste Mini-Übung")
    add_brand_h2(doc, mu["name"])
    add_brand_inline_caps_label(doc, "WANN EINSETZEN")
    add_brand_italic_note(doc, mu["wann_einsetzen"])
    add_brand_paragraph(doc, "")
    add_brand_inline_caps_label(doc, "DIE 6 SCHRITTE")
    for schritt in mu["schritte"]:
        # add_brand_number generiert seine eigene Word-Nummerierung — KEINE Zusatz-Nummer davor
        add_brand_number(doc, schritt)
    add_brand_paragraph(doc, "")
    add_brand_inline_caps_label(doc, "ZIEL")
    add_brand_italic_note(doc, mu["ziel"])
    add_brand_ornament(doc)
    add_brand_pagebreak(doc)

    # ---------- TEIL 7 — BRIDGE (stärker führend, 3-teilig) ----------
    br = a["teil_7_bridge"]
    add_brand_h1(doc, "Wie es weitergehen kann")

    # Teil A: Verstehen reicht nicht
    text_a = br.get("verstehen_reicht_nicht") or br.get("validierung", "")
    for absatz in _absaetze(text_a):
        add_brand_paragraph(doc, absatz)
    add_brand_paragraph(doc, "")
    add_brand_ornament(doc)
    add_brand_paragraph(doc, "")

    # Teil B: Nächster Schritt (klar führend)
    text_b = br.get("naechster_schritt") or br.get("soft_bridge", "")
    for absatz in _absaetze(text_b):
        add_brand_paragraph(doc, absatz)
    add_brand_paragraph(doc, "")

    # Teil C: Soft-Einladung
    text_c = br.get("soft_einladung")
    if text_c:
        add_brand_italic_note(doc, text_c)
        add_brand_paragraph(doc, "")

    # ---------- SCHLUSS ----------
    add_brand_ornament(doc)
    add_brand_paragraph(doc, "")
    for absatz in _absaetze(a["schluss"]):
        add_brand_paragraph(doc, absatz)

    # Header (Logo) + Footer (Website/Mail/Seite) auf allen Seiten außer Cover
    _add_header_footer(doc)

    save_branded_doc(doc, str(output_path))
    print(f"  ✓ DOCX gespeichert: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        # Default: alle drei Auswertungen rendern
        json_files = sorted((PROJECT_ROOT / "output").glob("profil_*_auswertung.json"))
    else:
        json_files = [Path(p) for p in sys.argv[1:]]

    for json_path in json_files:
        name_kurz = json_path.stem.replace("_auswertung", "")
        output_path = PROJECT_ROOT / "output" / f"Emotional-Stability-Analyse_{name_kurz}.docx"
        print(f"\n→ {json_path.name}")
        render_auswertung_pdf(json_path, output_path)
