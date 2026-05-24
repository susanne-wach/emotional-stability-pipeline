"""
Your Balance — branded Word-Dokument-Builder.

Dieses Modul stellt sicher, dass JEDES generierte Word-Dokument für Susanne Wachter /
Your Balance konsistent dem Master-Styleguide entspricht.

DOKUMENTEN-BRAND-SPEZIFIKATION (Master-Styleguide, Stand 2026-05, neue CI):
  - Schrift: Georgia durchgängig (universell installiert, Pages-/Word-kompatibel)
  - Cover-Caps-Label: 10pt regular CAPS, Gold #CFAC5E       (z. B. „KONTEXT-DOKUMENT")
  - Cover-Title:      26pt bold, Magenta #DD3089            (z. B. „Branding Guide")
  - Cover-Subtitle:   14pt regular, Aubergine #5B2A4A
  - Cover-Attribution:11pt regular, Aubergine #5B2A4A
  - Cover-URL:        11pt regular, Magenta #DD3089
  - Cover-Datum:      10pt regular, Grau #888888
  - H1 (Sektion):     18pt bold, Magenta #DD3089
  - H2:               14pt bold, Magenta #DD3089
  - Caps-Label inline:11pt bold CAPS, Gold #CFAC5E          (z. B. „PRIMÄRE BRANDFARBEN:")
  - Body:             11pt regular, Aubergine #5B2A4A        (NICHT reines Schwarz!)
  - Italic-Hinweis:   11pt italic, Aubergine #5B2A4A
  - Sekundär/Footer:  10pt regular, Grau #888888
  - Akzent:           Magenta #DD3089 — URL, Tabellen-Header, Callout-Label
  - Page-Margins:     2,54 cm (1 inch) rundum
  - Hintergrund:      Standard-Weiß (kein Softcream — das ist Web!)
  - Ornament:         ── ✦ ── (zentriert, Gold)

WICHTIG — Web- und Dokumenten-Anwendung nicht verwechseln:
  - WEB (Homepage your-balance.at): Softcream + Cormorant Garamond Light + Montserrat
  - DOKUMENTE (.docx): DIESES Modul = Georgia + weiß + Aubergine/Gold/Magenta

USAGE:
    import sys
    sys.path.insert(0, "/Users/susannewachter/Documents/07_Claude/resources/brand")
    from your_balance_docx import (
        create_branded_doc, BRAND,
        add_brand_cover,
        add_brand_caps_label, add_brand_title, add_brand_subtitle,
        add_brand_attribution, add_brand_url, add_brand_date,
        add_brand_h1, add_brand_h2,
        add_brand_inline_caps_label, add_brand_paragraph, add_brand_italic_note,
        add_brand_bullet, add_brand_number, add_brand_quote,
        add_brand_ornament, add_brand_pagebreak,
        add_brand_table, add_brand_callout, add_brand_mono,
        save_branded_doc,
    )

    doc = create_branded_doc()
    add_brand_cover(
        doc,
        caps_label="KONTEXT-DOKUMENT",
        title="Mein Titel",
        subtitle="Untertitel mit Beschreibung",
        attribution="Susanne Wachter – Your Balance",
        url="your-balance.at",
        date_text="Stand: Mai 2026",
    )
    add_brand_pagebreak(doc)
    add_brand_h1(doc, "Erste Sektion")
    add_brand_inline_caps_label(doc, "WICHTIG:")
    add_brand_paragraph(doc, "Body-Text in Dunkelbraun-Lila…")
    save_branded_doc(doc, "/path/to/output.docx")

Quelle (verbindlich):
  /Users/susannewachter/Documents/07_Claude/context/04-branding-guide.md  (Master-Styleguide)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# BRAND-KONSTANTEN — direkt aus 04-Branding-Guide.docx gemessen
# ============================================================
class BRAND:
    """Your-Balance-Dokumenten-Brand (Quelle: context/04-branding-guide.md, Master-Styleguide)."""

    # Schrift
    FONT = "Georgia"
    FONT_MONO = "Consolas"

    # Farben (Hex ohne #) — neue CI ab Mai 2026
    AUBERGINE = "5B2A4A"   # Hauptfarbe — Body, Subtitle, Attribution
    MAGENTA   = "DD3089"   # Cover-Title, H1, H2, URL, Tabellen-Header-Hintergrund, Callout-Label
    GOLD      = "CFAC5E"   # Caps-Labels (Cover & inline), Ornament
    GREY      = "888888"   # funktionales Neutral — Sekundär, Datum, Footer, Code
    WHITE     = "FFFFFF"   # Header-Text in Tabellen

    # Semantische Aliase
    BODY    = "5B2A4A"     # = AUBERGINE
    HEADING = "DD3089"     # = MAGENTA — Cover-Title, H1, H2

    # Legacy-Aliase für Rückwärtskompatibilität (alte CI)
    PINK       = "DD3089"  # historischer Name — mappt auf MAGENTA
    BLACK      = "5B2A4A"  # mappt auf AUBERGINE — kein reines Schwarz im Brand
    BOLD_BLACK = "5B2A4A"

    # Größen (Pt)
    SIZE_COVER_LABEL = 10        # Caps-Label im Cover (z. B. „KONTEXT-DOKUMENT")
    SIZE_TITLE = 26              # Cover-Title (z. B. „Branding Guide")
    SIZE_SUBTITLE = 14           # Cover-Subtitle
    SIZE_ATTRIBUTION = 11        # „Susanne Wachter – Your Balance"
    SIZE_URL = 11                # „your-balance.at"
    SIZE_DATE = 10               # „Stand: …"
    SIZE_H1 = 18                 # Sektion-H1
    SIZE_H2 = 14                 # Sektion-H2
    SIZE_BODY = 11               # Body & Italic-Hinweis
    SIZE_LABEL = 11              # Inline-Caps-Label (Gold, bold)
    SIZE_QUOTE = 10              # Original-Zitate
    SIZE_FOOTER = 10             # Sekundär-/Footer-Text

    # Layout
    MARGIN_CM = 2.54             # 1 inch — wie im Master


# ============================================================
# DOKUMENT-INITIALISIERUNG
# ============================================================
def create_branded_doc():
    """Erstellt ein neues Word-Dokument mit Georgia als Default-Schrift."""
    doc = Document()

    # Page-Setup: 2,54 cm Ränder rundum (wie im Master)
    for section in doc.sections:
        section.left_margin = Cm(BRAND.MARGIN_CM)
        section.right_margin = Cm(BRAND.MARGIN_CM)
        section.top_margin = Cm(BRAND.MARGIN_CM)
        section.bottom_margin = Cm(BRAND.MARGIN_CM)

    # Default-Style: Georgia 11pt Dunkelbraun-Lila
    normal = doc.styles["Normal"]
    normal.font.name = BRAND.FONT
    normal.font.size = Pt(BRAND.SIZE_BODY)
    normal.font.color.rgb = RGBColor.from_string(BRAND.BODY)
    _set_style_font(normal.element, BRAND.FONT)

    return doc


def _set_style_font(element, font_name):
    """Setzt rFonts für einen Style, damit Georgia auch in CJK/CS-Fallback greift."""
    try:
        rPr = element.get_or_add_rPr()
    except AttributeError:
        return
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)


def _styled_run(p, text, *, size=None, color=None, bold=False, italic=False, font=None):
    """Helper: Run mit konsistentem Brand-Styling anlegen."""
    run = p.add_run(text)
    run.font.name = font or BRAND.FONT
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    # rFonts auch im Run-Element setzen
    rPr = run.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font or BRAND.FONT)
    rFonts.set(qn("w:hAnsi"), font or BRAND.FONT)
    rFonts.set(qn("w:cs"), font or BRAND.FONT)
    rFonts.set(qn("w:eastAsia"), font or BRAND.FONT)
    return run


# ============================================================
# COVER — exakt nach Master-Aufbau
# ============================================================
def add_brand_caps_label(doc, text, *, align_center=True):
    """Cover-Caps-Label — 10pt regular CAPS, Gold. Wie „KONTEXT-DOKUMENT" im Master."""
    p = doc.add_paragraph()
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(p, text.upper(), size=BRAND.SIZE_COVER_LABEL, color=BRAND.GOLD)
    return p


def add_brand_title(doc, text, *, align_center=True):
    """Cover-Title — 26pt bold, Magenta."""
    p = doc.add_paragraph()
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(p, text, size=BRAND.SIZE_TITLE, color=BRAND.HEADING, bold=True)
    return p


def add_brand_subtitle(doc, text, *, align_center=True):
    """Cover-Subtitle — 14pt regular, Dunkelbraun-Lila. Wie „Farben, Typografie & …" im Master."""
    p = doc.add_paragraph()
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(p, text, size=BRAND.SIZE_SUBTITLE, color=BRAND.BODY)
    return p


def add_brand_attribution(doc, text, *, align_center=True):
    """Cover-Attribution — 11pt regular, Dunkelbraun-Lila. Wie „Susanne Wachter – Your Balance"."""
    p = doc.add_paragraph()
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(p, text, size=BRAND.SIZE_ATTRIBUTION, color=BRAND.BODY)
    return p


def add_brand_url(doc, text, *, align_center=True):
    """Cover-URL — 11pt regular, Magenta. Wie „your-balance.at"."""
    p = doc.add_paragraph()
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(p, text, size=BRAND.SIZE_URL, color=BRAND.MAGENTA)
    return p


def add_brand_date(doc, text, *, align_center=True):
    """Cover-Datum — 10pt regular, Grau. Wie „Stand: April 2026"."""
    p = doc.add_paragraph()
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(p, text, size=BRAND.SIZE_DATE, color=BRAND.GREY)
    return p


def add_brand_cover(doc, *, caps_label=None, title, subtitle=None,
                    attribution=None, url=None, date_text=None,
                    top_spacer_lines=10):
    """Erstellt einen kompletten Cover-Block exakt nach Master-Vorlage.

    Reihenfolge: vertikaler Spacer → Caps-Label → Title → Subtitle → leere Zeile
                 → Attribution → URL → Date → Page-Break
    """
    # Vertikaler Spacer oben (wie im Master, der Title sitzt nach langem Weißraum)
    for _ in range(top_spacer_lines):
        doc.add_paragraph()

    if caps_label:
        add_brand_caps_label(doc, caps_label)
    add_brand_title(doc, title)
    if subtitle:
        add_brand_subtitle(doc, subtitle)
    # Leere Zeile vor Attribution-Block
    doc.add_paragraph()
    if attribution:
        add_brand_attribution(doc, attribution)
    if url:
        add_brand_url(doc, url)
    if date_text:
        add_brand_date(doc, date_text)


# ============================================================
# SEKTIONEN / ÜBERSCHRIFTEN
# ============================================================
def add_brand_h1(doc, text):
    """H1 — 18pt bold, Magenta. Sektions-Überschrift."""
    p = doc.add_paragraph()
    _styled_run(p, text, size=BRAND.SIZE_H1, color=BRAND.HEADING, bold=True)
    return p


def add_brand_h2(doc, text):
    """H2 — 14pt bold, Magenta."""
    p = doc.add_paragraph()
    _styled_run(p, text, size=BRAND.SIZE_H2, color=BRAND.HEADING, bold=True)
    return p


def add_brand_inline_caps_label(doc, text):
    """Inline-Caps-Label — 11pt bold CAPS, Gold. Wie „PRIMÄRE HINTERGRUNDFARBEN:" im Master."""
    p = doc.add_paragraph()
    _styled_run(p, text.upper(), size=BRAND.SIZE_LABEL, color=BRAND.GOLD, bold=True)
    return p


# Legacy-Alias für altes API
def add_brand_h3_gold(doc, text):
    """[Legacy] Wie inline-caps-label, aber kein Auto-Uppercase."""
    p = doc.add_paragraph()
    _styled_run(p, text, size=BRAND.SIZE_LABEL, color=BRAND.GOLD, bold=True)
    return p


def add_brand_sub_label(doc, text):
    """[Legacy] Inline-Sub-Label — 11pt bold, Pink."""
    p = doc.add_paragraph()
    _styled_run(p, text, size=BRAND.SIZE_LABEL, color=BRAND.PINK, bold=True)
    return p


def add_brand_yb_label_paragraph(doc, label_text, body_text):
    """[Legacy] Kombi: Gold-bold Label + Body-Text."""
    p = doc.add_paragraph()
    _styled_run(p, label_text, size=BRAND.SIZE_LABEL, color=BRAND.GOLD, bold=True)
    _styled_run(p, " " + body_text, size=BRAND.SIZE_BODY, color=BRAND.BODY)
    return p


# ============================================================
# BODY / ABSÄTZE
# ============================================================
def add_brand_paragraph(doc, text="", *, bold=False, italic=False, color=None, size=None):
    """Standard-Fließtext-Absatz — 11pt Dunkelbraun-Lila."""
    p = doc.add_paragraph()
    _styled_run(p, text,
                size=size or BRAND.SIZE_BODY,
                color=color or BRAND.BODY,
                bold=bold, italic=italic)
    return p


def add_brand_italic_note(doc, text):
    """Italic-Hinweis — 11pt italic, Dunkelbraun-Lila. Wie „Farbregel: …" im Master."""
    p = doc.add_paragraph()
    _styled_run(p, text, size=BRAND.SIZE_BODY, color=BRAND.BODY, italic=True)
    return p


def add_brand_bullet(doc, text):
    """Bullet-List-Item im Brand-Style — Georgia 11pt Dunkelbraun-Lila."""
    p = doc.add_paragraph(style="List Bullet")
    for run in list(p.runs):
        run.text = ""
    _styled_run(p, text, size=BRAND.SIZE_BODY, color=BRAND.BODY)
    return p


def add_brand_number(doc, text):
    """Numbered-List-Item im Brand-Style."""
    p = doc.add_paragraph(style="List Number")
    for run in list(p.runs):
        run.text = ""
    _styled_run(p, text, size=BRAND.SIZE_BODY, color=BRAND.BODY)
    return p


def add_brand_quote(doc, text, label_prefix=""):
    """Original-Quote — 10pt italic, Grau."""
    p = doc.add_paragraph()
    full = (label_prefix + text) if label_prefix else text
    _styled_run(p, full, size=BRAND.SIZE_QUOTE, color=BRAND.GREY, italic=True)
    return p


def add_brand_lead_quote(doc, text):
    """[Legacy] Lead-Zitat unter dem Titel — italic, Body-Farbe."""
    p = doc.add_paragraph()
    _styled_run(p, text, size=BRAND.SIZE_BODY, color=BRAND.BODY, italic=True)
    return p


def add_brand_mono(doc, text):
    """Monospace-Block für Code/IDs — Consolas 9pt Grau."""
    p = doc.add_paragraph()
    _styled_run(p, text, size=9, color=BRAND.GREY, font=BRAND.FONT_MONO)
    return p


# ============================================================
# DESIGN-ELEMENTE
# ============================================================
def add_brand_ornament(doc):
    """Markenornament `── ✦ ──` zentriert in Gold. Trennt Sektionen elegant."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(p, "── ✦ ──", size=BRAND.SIZE_BODY, color=BRAND.GOLD)
    return p


def add_brand_pagebreak(doc):
    """Seitenumbruch."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    return p


# ============================================================
# TABELLEN
# ============================================================
def add_brand_table(doc, headers, rows, col_widths_cm=None):
    """Tabelle mit Magenta-Header (Background, weiß bold) + Body in Aubergine."""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header-Zeile mit Pink-Hintergrund
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), BRAND.MAGENTA)
        tc_pr.append(shd)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = BRAND.FONT
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(BRAND.WHITE)
        # rFonts auch hier setzen
        rPr = run.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), BRAND.FONT)
        rFonts.set(qn("w:hAnsi"), BRAND.FONT)

    # Data-Zeilen
    for r_idx, row_data in enumerate(rows, start=1):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            header_lower = headers[c_idx].lower() if c_idx < len(headers) else ""
            is_mono = any(k in header_lower for k in ["id", "page", "code", "hex", "suffix"])
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(BRAND.BODY)
            run.font.name = BRAND.FONT_MONO if is_mono else BRAND.FONT
            rPr = run.element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            font_used = BRAND.FONT_MONO if is_mono else BRAND.FONT
            rFonts.set(qn("w:ascii"), font_used)
            rFonts.set(qn("w:hAnsi"), font_used)

    if col_widths_cm:
        for c_idx, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[c_idx].width = Cm(w)

    return table


def add_brand_callout(doc, text, label="HINWEIS"):
    """Highlight-Box als 1×1-Tabelle mit Magenta-Label + Body-Text."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side in ["top", "bottom", "left", "right"]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), "120")
        m.set(qn("w:type"), "dxa")
        tc_mar.append(m)
    tc_pr.append(tc_mar)

    p = cell.paragraphs[0]
    _styled_run(p, label.upper() + ": ", size=BRAND.SIZE_BODY, color=BRAND.MAGENTA, bold=True)
    _styled_run(p, text, size=BRAND.SIZE_BODY, color=BRAND.BODY)
    return table


# ============================================================
# SAVE & SELF-CHECK
# ============================================================
def save_branded_doc(doc, path, run_brand_check=True):
    """Speichert das Dokument und führt einen Brand-Self-Check aus."""
    doc.save(path)
    import os
    size_kb = os.path.getsize(path) / 1024
    print(f"✓ Saved: {path}  ({size_kb:.1f} KB)")
    if run_brand_check:
        _brand_self_check(doc)


def _brand_self_check(doc):
    """Selbst-Check gegen den Master-Styleguide (neue CI, Stand Mai 2026)."""
    print("Brand-Self-Check (Quelle: context/04-branding-guide.md):")
    has_georgia = False
    has_aubergine = False
    has_gold = False
    has_magenta = False
    has_calibri_violation = False
    has_pure_black_violation = False
    has_old_ci_violation = False

    def _scan(paragraphs):
        nonlocal has_georgia, has_aubergine, has_gold, has_magenta
        nonlocal has_calibri_violation, has_pure_black_violation, has_old_ci_violation
        for p in paragraphs:
            for run in p.runs:
                if run.font.name == BRAND.FONT:
                    has_georgia = True
                if run.font.name == "Calibri":
                    has_calibri_violation = True
                if run.font.color and run.font.color.rgb:
                    hex_color = str(run.font.color.rgb).upper()
                    if hex_color == BRAND.AUBERGINE.upper(): has_aubergine = True
                    if hex_color == BRAND.GOLD.upper(): has_gold = True
                    if hex_color == BRAND.MAGENTA.upper(): has_magenta = True
                    if hex_color == "000000": has_pure_black_violation = True
                    # alte CI: Dunkelbraun-Lila #3A2E3D / Web-Gold #C9A84C
                    if hex_color in ("3A2E3D", "C9A84C", "1C1208", "2A1F0E"):
                        has_old_ci_violation = True

    _scan(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _scan(cell.paragraphs)

    print(f"  Georgia verwendet:                   {'✓' if has_georgia else '⚠ NEIN'}")
    print(f"  Magenta #DD3089 (Überschriften):     {'✓' if has_magenta else '⚠ NEIN'}")
    print(f"  Aubergine #5B2A4A (Body-Text):       {'✓' if has_aubergine else '⚠ NEIN'}")
    print(f"  Gold #CFAC5E (Caps-Labels):          {'✓' if has_gold else '⚠ NEIN'}")
    if has_calibri_violation:
        print(f"  ⚠ FEHLER: Calibri verwendet — gehört NICHT in den Dokumenten-Brand!")
    if has_pure_black_violation:
        print(f"  ⚠ FEHLER: Reines Schwarz #000000 — Text muss Aubergine #5B2A4A sein!")
    if has_old_ci_violation:
        print(f"  ⚠ FEHLER: Alte CI-Farbe gefunden (#3A2E3D/#C9A84C/#1C1208) — bitte neue CI verwenden!")
